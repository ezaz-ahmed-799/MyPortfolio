from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Ezaz_Ahmed_Resume.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
BORDER = "D9E2EC"
FILL = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)

    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def set_font(run, size=10.5, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, size=10.5, color=None, bold=False, italic=False, after=4, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, size=11.5, color=BLUE, bold=True)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "D9E2EC")
    border.append(bottom)
    p_pr.append(border)
    return p


def add_bullet(doc, text, after=2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size=10.2)
    return p


def add_role(doc, title, org, detail):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    set_font(r, size=10.8, color=NAVY, bold=True)
    r = p.add_run(f" | {org}")
    set_font(r, size=10.8, color=DARK_BLUE, bold=True)
    r = p.add_run(f" | {detail}")
    set_font(r, size=9.5, color=MUTED)


def add_project(doc, title, label, stack, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    set_font(r, size=10.6, color=NAVY, bold=True)
    r = p.add_run(f" - {label}")
    set_font(r, size=10.2, color=MUTED, italic=True)
    add_bullet(doc, f"{description} Tech: {stack}", after=2)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.2)
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.15


def build_resume():
    doc = Document()
    configure_styles(doc)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run("EZAZ AHMED")
    set_font(run, size=22, color=NAVY, bold=True)

    tagline = add_para(
        doc,
        "Full Stack Developer | AI Researcher | MERN, Spring Boot, Machine Learning",
        size=10.4,
        color=MUTED,
        bold=True,
        after=3,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    tagline.paragraph_format.line_spacing = 1.0

    contact = add_para(
        doc,
        "ezazahmedmd555@gmail.com | github.com/ezaz-ahmed-799 | linkedin.com/in/ezaz-ahmed-775898277 | ORCID: 0009-0009-6572-8661",
        size=9.3,
        color=MUTED,
        after=6,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    contact.paragraph_format.line_spacing = 1.0

    add_heading(doc, "Professional Summary")
    add_para(
        doc,
        "Full Stack Developer and AI Researcher from Andhra Pradesh, India, specializing in MERN Stack, Java Spring Boot, and machine learning systems. Experienced in building scalable user-centric applications and contributing to published research in computer vision, space weather prediction, cloud security, and intelligent architectures.",
        size=10.4,
        after=4,
    )

    add_heading(doc, "Technical Skills")
    skills = doc.add_table(rows=4, cols=2)
    set_table_width(skills, [1.8, 5.2])
    rows = [
        ("Frontend", "React, TypeScript, Angular, responsive UI development"),
        ("Backend", "Node.js, Express, Java, Spring Boot, REST APIs, JWT"),
        ("Databases", "MongoDB, PostgreSQL"),
        ("AI / Research", "Python, TensorFlow, CNN, 3D CNN, Deep Learning, Machine Learning, Computer Vision, predictive modeling"),
    ]
    for row, (label, value) in zip(skills.rows, rows):
        set_cell_shading(row.cells[0], FILL)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=9.8, color=NAVY, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.8)

    add_heading(doc, "Experience")
    add_role(doc, "Full Stack Developer", "VCR Tech Solutions", "March 2026 - Present")
    add_bullet(doc, "Build scalable full stack applications using MERN Stack and Java Spring Boot with a focus on performance, architecture, and user-centric systems.")
    add_bullet(doc, "Contribute across frontend, backend, database, and integration layers to deliver maintainable production-ready software.")

    add_role(doc, "Program Analyst Trainee", "Cognizant", "June 2025 - February 2026")
    add_bullet(doc, "Worked on enterprise applications using Angular, Spring Boot, MongoDB, and PostgreSQL.")
    add_bullet(doc, "Contributed to scalable backend systems and modern frontend experiences in a collaborative engineering environment.")

    add_role(doc, "Program Analyst Trainee Intern", "Cognizant", "January 2025 - June 2025")
    add_bullet(doc, "Completed intensive enterprise training in React, Java Spring Boot, and modern full stack development practices.")
    add_bullet(doc, "Built professional coding habits around standards, collaboration workflows, and industry-grade software engineering methodology.")

    add_heading(doc, "Selected Projects")
    add_project(
        doc,
        "Human Action Recognition",
        "AI Research Project",
        "Python, CNN, Deep Learning, Computer Vision, TensorFlow",
        "Developed a computer vision system for video sequence classification and intelligent activity detection using CNN architectures.",
    )
    add_project(
        doc,
        "Hospital Management System",
        "Enterprise Full Stack System",
        "React, Spring Boot, JWT, PostgreSQL, REST APIs",
        "Built a healthcare management platform with authentication, appointment scheduling, patient workflows, and role-based access control.",
    )
    add_project(
        doc,
        "Solar Flare Prediction",
        "Predictive AI System",
        "Machine Learning, Python, Predictive Modeling, Data Mining, Neural Networks",
        "Created a deep learning-driven forecasting system for space weather analysis using scientific datasets and comparative neural architectures.",
    )
    add_project(
        doc,
        "Student Platform",
        "Full Stack Platform",
        "React, Node.js, MongoDB, Express, REST APIs",
        "Designed an interactive platform for academic workflows, collaboration, dashboard management, and scalable user experiences.",
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    add_heading(doc, "Research Publications")
    add_role(doc, "Exploring CNN-Based Algorithms for Human Action Recognition in Videos", "Springer Nature / BROADNETS 2024", "Published, Scopus Indexed")
    add_bullet(doc, "Evaluated Two-Stream CNN, CNN + LSTM, and 3D CNN architectures on HMDB-51; DOI: 10.1007/978-3-031-81171-5_11; Scopus EID: 2-s2.0-85219192326.")
    add_bullet(doc, "Found that 3D CNN delivered the strongest accuracy and computational efficiency through direct spatiotemporal representation learning.")

    add_role(doc, "Solar Flare Prediction Using Machine Learning and Deep Learning Techniques", "Scopus Indexed Research Publication", "Published")
    add_bullet(doc, "Studied machine learning and deep learning methods for forecasting solar flare activity from space weather datasets; Scopus EID: 2-s2.0-105034141640.")

    add_role(doc, "Machine Learning-Based Detection and Mitigation of Privilege Escalation Attacks", "Cloud Security Research", "Under Publication, January 2025")
    add_bullet(doc, "Proposed a behavioral anomaly detection model for cloud privilege escalation attacks, reporting 98.94% binary and 98.92% multiclass classification accuracy.")

    add_heading(doc, "Education")
    add_role(doc, "B.Tech in Computer Science and Engineering", "Seshadri Rao Gudlavalleru Engineering College", "2021 - 2025")
    add_bullet(doc, "Specialization in Artificial Intelligence and Machine Learning; Andhra Pradesh, India.")
    add_role(doc, "MS Computer Science Admit", "University of Louisville", "R1 Research Institution")
    add_bullet(doc, "Admitted with International Resident Tuition Grant recognition; did not attend due to personal reasons.")

    add_heading(doc, "Additional Strengths")
    add_bullet(doc, "Research output: 2 Scopus-indexed papers, 1 verified citation, h-index 1.")
    add_bullet(doc, "Personal strengths shaped by competitive sports, strategy-focused gaming, and story-rich creative media: discipline, persistence, teamwork, and analytical thinking.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("Ezaz Ahmed - Resume")
    set_font(r, size=8.5, color=MUTED)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_resume())
